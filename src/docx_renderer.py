import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from config import (
    DOCX_BULLET_SYMBOL,
    DOCX_DEFINITION_SEPARATOR,
    DOCX_EXAMPLES_HEADING,
    DOCX_FALLBACK_FONT,
    DOCX_FILE_EXTENSION,
    DOCX_HEADING_1_STYLE_NAME,
    DOCX_HEADING_2_STYLE_NAME,
    DOCX_HEADING_3_STYLE_NAME,
    DOCX_MAX_FONT_SIZE,
    DOCX_MAX_MARGIN,
    DOCX_MIN_FONT_SIZE,
    DOCX_MIN_MARGIN,
    DOCX_NESTED_BULLET_SYMBOL,
    DOCX_NORMAL_STYLE_NAME,
    DOCX_SUBTITLE_STYLE_NAME,
    DOCX_SUMMARY_HEADING,
    DOCX_TIMESTAMP_FORMAT,
    DOCX_TITLE_STYLE_NAME,
    GENERATED_DOCUMENT_FILE_PREFIX,
    GENERATED_DOCUMENTS_DIR,
)

logger = logging.getLogger(__name__)

class DocxRenderer:
    """Creates formatted Word documents from structured lecture-note data."""

    REQUIRED_PROFILE_SECTIONS = (
        "page",
        "body",
        "title",
        "heading_1",
        "heading_2",
        "heading_3",
        "bullet",
    )

    def __init__(self, output_dir: Path = GENERATED_DOCUMENTS_DIR) -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def render(
        self,
        notes: dict[str, Any],
        style_profile: dict[str, Any],
        filename: str | None = None,
    ) -> Path:
        """Create and save a fully formatted DOCX document."""

        notes = self._validate_notes(notes)
        profile = self._validate_profile(style_profile)
        output_path = self._build_output_path(filename)

        document = Document()
        self._configure_page(document, profile["page"])
        self._configure_document_styles(document, profile)

        self._add_title(document, notes["title"], profile["title"])

        if notes.get("subtitle"):
            self._add_subtitle(document, notes["subtitle"], profile["body"])

        self._add_sections(document, notes["sections"], profile)

        if notes.get("summary"):
            self._add_summary(document, notes["summary"], profile)

        try:
            document.save(output_path)
        except (OSError, PermissionError) as error:
            logger.exception("Could not save DOCX document (file=%s).", output_path.name)
            raise RuntimeError(
                "The formatted Word document could not be saved. "
                "Make sure the file is not already open in Word."
            ) from error

        logger.info("Formatted DOCX document saved (file=%s).", output_path.name)
        return output_path

    def _configure_page(self, document: DocumentObject, page: dict[str, Any]) -> None:
        """Apply detected page dimensions and margins."""

        width = self._safe_number(page.get("width"), 595.3, 300.0, 1000.0)
        height = self._safe_number(page.get("height"), 841.9, 400.0, 1500.0)

        margins = {
            "left_margin": self._safe_number(
                page.get("margin_left"), 72.0, DOCX_MIN_MARGIN, DOCX_MAX_MARGIN
            ),
            "right_margin": self._safe_number(
                page.get("margin_right"), 72.0, DOCX_MIN_MARGIN, DOCX_MAX_MARGIN
            ),
            "top_margin": self._safe_number(
                page.get("margin_top"), 72.0, DOCX_MIN_MARGIN, DOCX_MAX_MARGIN
            ),
            "bottom_margin": self._safe_number(
                page.get("margin_bottom"), 72.0, DOCX_MIN_MARGIN, DOCX_MAX_MARGIN
            ),
        }

        for section in document.sections:
            section.page_width = Pt(width)
            section.page_height = Pt(height)

            for attribute, value in margins.items():
                setattr(section, attribute, Pt(value))

    def _configure_document_styles(self, document: DocumentObject, profile: dict[str, Any],) -> None:
        """Configure Word's built-in paragraph styles."""

        self._configure_style(
            document,
            DOCX_NORMAL_STYLE_NAME,
            profile["body"],
            use_line_spacing=True,
        )
        self._configure_style(document, DOCX_TITLE_STYLE_NAME, profile["title"])

        subtitle_style = {
            **profile["body"],
            "italic": True,
            "font_size": max(
                float(profile["body"]["font_size"]),
                float(profile["heading_3"]["font_size"]),
            ),
        }

        self._configure_style(document, DOCX_SUBTITLE_STYLE_NAME, subtitle_style)
        self._configure_style(
            document, DOCX_HEADING_1_STYLE_NAME, profile["heading_1"]
        )
        self._configure_style(
            document, DOCX_HEADING_2_STYLE_NAME, profile["heading_2"]
        )
        self._configure_style(
            document, DOCX_HEADING_3_STYLE_NAME, profile["heading_3"]
        )

    def _configure_style(
        self,
        document: DocumentObject,
        style_name: str,
        style_data: dict[str, Any],
        use_line_spacing: bool = False,
    ) -> None:
        """Apply text and paragraph properties to a Word style."""

        try:
            word_style = document.styles[style_name]
        except KeyError as error:
            raise ValueError(f"Word style {style_name!r} is unavailable.") from error

        font_name = self._safe_font(style_data.get("font"))
        font_size = self._font_size(style_data.get("font_size"), 11.0)

        word_style.font.name = font_name
        word_style.font.size = Pt(font_size)
        word_style.font.bold = self._safe_bool(style_data.get("bold"), False)
        word_style.font.italic = self._safe_bool(style_data.get("italic"), False)
        word_style.font.color.rgb = self._parse_color(
            style_data.get("color", "#000000")
        )

        self._set_font_family(word_style.element.rPr, font_name)

        paragraph_format = word_style.paragraph_format
        paragraph_format.space_before = Pt(
            self._get_spacing(
                style_data,
                "spacing_before",
                "paragraph_spacing_before",
            )
        )
        paragraph_format.space_after = Pt(
            self._get_spacing(
                style_data,
                "spacing_after",
                "paragraph_spacing_after",
            )
        )

        if use_line_spacing:
            paragraph_format.line_spacing = Pt(
                self._safe_number(
                    style_data.get("line_spacing_points"),
                    font_size * 1.15,
                    font_size,
                    60.0,
                )
            )

    def _add_title(
        self,
        document: DocumentObject,
        text: str,
        style: dict[str, Any],
    ) -> None:
        paragraph = document.add_paragraph(style=DOCX_TITLE_STYLE_NAME)
        paragraph.alignment = self._parse_alignment(style.get("alignment", "left"))
        paragraph.paragraph_format.space_after = Pt(
            self._safe_number(style.get("spacing_after"), 12.0, 0.0, 72.0)
        )

        run = paragraph.add_run(text.strip())
        self._apply_run_style(run, style)

    def _add_subtitle(
        self,
        document: DocumentObject,
        text: str,
        body_style: dict[str, Any],
    ) -> None:
        paragraph = document.add_paragraph(style=DOCX_SUBTITLE_STYLE_NAME)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.add_run(text.strip())
        self._apply_run_style(run, {**body_style, "italic": True})

    def _add_sections(
        self,
        document: DocumentObject,
        sections: list[dict[str, Any]],
        profile: dict[str, Any],
    ) -> None:
        """Render every generated note section and its contents."""

        for index, section in enumerate(sections, start=1):
            heading = str(section.get("heading", "")).strip()

            if not heading:
                logger.warning("Skipping section %d because its heading is empty.", index)
                continue

            level = self._get_heading_level(section)
            self._add_heading(document, heading, level, profile)

            for paragraph in section.get("paragraphs", []):
                if isinstance(paragraph, str) and paragraph.strip():
                    self._add_body_paragraph(document, paragraph, profile["body"])

            for bullet in section.get("bullets", []):
                if isinstance(bullet, dict):
                    self._add_bullet(document, bullet, profile)

            definitions = section.get("definitions", [])

            if definitions:
                self._add_definitions(document, definitions, profile)

            examples = section.get("examples", [])

            if examples:
                self._add_examples(document, examples, profile)

    def _add_heading(
        self,
        document: DocumentObject,
        text: str,
        level: int,
        profile: dict[str, Any],
    ) -> None:
        style_name = self._heading_style_name(level)
        style_data = profile[f"heading_{level}"]

        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.space_before = Pt(
            self._safe_number(style_data.get("spacing_before"), 8.0, 0.0, 72.0)
        )
        paragraph.paragraph_format.space_after = Pt(
            self._safe_number(style_data.get("spacing_after"), 4.0, 0.0, 72.0)
        )

        run = paragraph.add_run(self._format_heading_text(text, level))
        self._apply_run_style(run, style_data)

    def _add_body_paragraph(
        self,
        document: DocumentObject,
        text: str,
        body_style: dict[str, Any],
        italic: bool | None = None,
        left_indent: float = 0.0,
    ) -> None:
        paragraph = document.add_paragraph(style=DOCX_NORMAL_STYLE_NAME)
        self._apply_body_paragraph_format(
            paragraph,
            body_style,
            left_indent=left_indent,
        )

        run_style = body_style.copy()

        if italic is not None:
            run_style["italic"] = italic

        run = paragraph.add_run(text.strip())
        self._apply_run_style(run, run_style)

    def _add_bullet(
        self,
        document: DocumentObject,
        bullet: dict[str, Any],
        profile: dict[str, Any],
    ) -> None:
        """Render one main bullet and all its nested children."""

        text = bullet.get("text", "")

        if not isinstance(text, str) or not text.strip():
            return

        bullet_style = profile["bullet"]
        body_style = profile["body"]

        indent = self._safe_number(
            bullet_style.get("indent_left"),
            24.0,
            0.0,
            200.0,
        )
        hanging = self._safe_number(
            bullet_style.get("hanging_indent"),
            12.0,
            0.0,
            indent,
        )
        nested_increment = self._safe_number(
            bullet_style.get("nested_indent_increment"),
            indent,
            0.0,
            200.0,
        )

        self._add_bullet_paragraph(
            document=document,
            text=text,
            symbol=DOCX_BULLET_SYMBOL,
            level_indent=indent,
            hanging_indent=hanging,
            style=bullet_style,
            body_style=body_style,
        )

        for child in bullet.get("children", []):
            if not isinstance(child, str) or not child.strip():
                continue

            self._add_bullet_paragraph(
                document=document,
                text=child,
                symbol=DOCX_NESTED_BULLET_SYMBOL,
                level_indent=indent + nested_increment,
                hanging_indent=hanging,
                style=bullet_style,
                body_style=body_style,
            )

    def _add_bullet_paragraph(
        self,
        document: DocumentObject,
        text: str,
        symbol: str,
        level_indent: float,
        hanging_indent: float,
        style: dict[str, Any],
        body_style: dict[str, Any],
    ) -> None:
        paragraph = document.add_paragraph(style=DOCX_NORMAL_STYLE_NAME)
        paragraph.paragraph_format.left_indent = Pt(level_indent)

        # A negative first-line indent creates a hanging indent.
        paragraph.paragraph_format.first_line_indent = Pt(-hanging_indent)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(
            self._safe_number(style.get("spacing_after"), 3.0, 0.0, 72.0)
        )
        paragraph.paragraph_format.line_spacing = Pt(
            self._body_line_spacing(body_style)
        )

        run = paragraph.add_run(f"{symbol} {text.strip()}")
        self._apply_run_style(run, {
            **body_style,
            "font": style.get("font", body_style.get("font")),
            "font_size": style.get("font_size", body_style.get("font_size")),
            "color": style.get("color", body_style.get("color")),
        })

    def _add_definitions(
        self,
        document: DocumentObject,
        definitions: list[Any],
        profile: dict[str, Any],
    ) -> None:
        """Render definitions in the user's inline note-taking style."""

        body_style = profile["body"]

        for definition in definitions:
            if not isinstance(definition, dict):
                continue

            term = definition.get("term", "")
            meaning = definition.get("definition", "")

            if not isinstance(term, str) or not isinstance(meaning, str):
                continue

            term = term.strip()
            meaning_text = meaning.strip()

            if not term or not meaning_text:
                continue

            paragraph = document.add_paragraph(
                style=DOCX_NORMAL_STYLE_NAME
            )

            self._apply_body_paragraph_format(
                paragraph,
                body_style,
            )

            term_run = paragraph.add_run(term)

            self._apply_run_style(
                term_run,
                {
                    **body_style,
                    "bold": True,
                },
            )

            lower_meaning = meaning_text.lower()

            existing_connectors = (
                "is ", "are ", "means ", "describes ", "refers to ", "represents ", "defines ",
                "explains ", "shows ", "demonstrates ", "indicates ", "occurs ", "happens ", "works ", "behaves ",
                "allows ", "enables ", "uses ", "takes ", "accepts ", "returns ", "produces ", "creates ",
                "generates ", "converts ", "transforms ", "stores ", "contains ", "provides ",
                "checks ", "tests ", "verifies ", "validates ", "ensures ", "identifies ", "detects ",
                "determines ", "measures ", "calculates ", "compares ", "evaluates ", "analyzes ",
                "consists of ", "includes ", "involves ", "connects ", "combines ", "organizes ", "groups ", "separates ",
                "helps ", "prevents ", "protects ", "controls ", "manages ", "supports ", "improves ",
                "reduces ", "increases ", "focuses on ", "depends on ",
            )

            if lower_meaning.startswith(existing_connectors):
                connector = " "
            else:
                connector = " describes "
                meaning_text = meaning_text[0].lower() + meaning_text[1:]

            meaning_run = paragraph.add_run(
                connector + meaning_text
            )

            self._apply_run_style(
                meaning_run,
                body_style,
            )

    def _add_examples(
        self,
        document: DocumentObject,
        examples: list[Any],
        profile: dict[str, Any],
    ) -> None:
        """Render examples as ordinary explanatory notes."""

        valid_examples = [
            example.strip()
            for example in examples
            if isinstance(example, str) and example.strip()
        ]

        for example in valid_examples:
            text = example

            if not text.lower().startswith(
                ("for example", "for instance", "e.g.")
            ):
                text = f"For example, {text[0].lower()}{text[1:]}"

            self._add_body_paragraph(
                document=document,
                text=text,
                body_style=profile["body"],
                italic=False,
                left_indent=0.0,
            )

    def _add_summary(
        self,
        document: DocumentObject,
        summary: list[Any],
        profile: dict[str, Any],
    ) -> None:
        """Render the final summary as a Heading 1 section with bullets."""

        valid_items = [
            item.strip()
            for item in summary
            if isinstance(item, str) and item.strip()
        ]

        if not valid_items:
            return

        self._add_heading(
            document,
            DOCX_SUMMARY_HEADING,
            level=2,
            profile=profile,
        )

        if len(valid_items) == 1:
            self._add_body_paragraph(
                document=document,
                text=valid_items[0],
                body_style=profile["body"],
            )
            return

        for item in valid_items:
            self._add_bullet(
                document=document,
                bullet={
                    "text": item,
                    "children": [],
                },
                profile=profile,
            )

    def _format_heading_text(self, text: str, level: int) -> str:
        """Format headings similarly to the reference notes."""

        text = text.strip()
        if not text:
            return ""

        if level == 1:
            return text.upper()
        return text

    def _apply_body_paragraph_format(
        self,
        paragraph: Any,
        body_style: dict[str, Any],
        left_indent: float = 0.0,
    ) -> None:
        paragraph.paragraph_format.left_indent = Pt(left_indent)
        paragraph.paragraph_format.space_before = Pt(
            self._safe_number(
                body_style.get("paragraph_spacing_before"),
                0.0,
                0.0,
                72.0,
            )
        )
        paragraph.paragraph_format.space_after = Pt(
            self._safe_number(
                body_style.get("paragraph_spacing_after"),
                6.0,
                0.0,
                72.0,
            )
        )
        paragraph.paragraph_format.line_spacing = Pt(
            self._body_line_spacing(body_style)
        )

    def _body_line_spacing(self, body_style: dict[str, Any]) -> float:
        font_size = self._font_size(body_style.get("font_size"), 11.0)

        return self._safe_number(
            body_style.get("line_spacing_points"),
            font_size * 1.15,
            font_size,
            60.0,
        )

    def _apply_run_style(self, run: Any, style_data: dict[str, Any]) -> None:
        font_name = self._safe_font(style_data.get("font"))
        font_size = self._font_size(style_data.get("font_size"), 11.0)

        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = self._safe_bool(style_data.get("bold"), False)
        run.italic = self._safe_bool(style_data.get("italic"), False)
        run.font.color.rgb = self._parse_color(
            style_data.get("color", "#000000")
        )
        self._set_font_family(run._element.rPr, font_name)

    def _set_font_family(self, run_properties: Any, font_name: str) -> None:
        """Apply one font family across common Word character sets."""

        if run_properties is None:
            return

        run_properties.rFonts.set(qn("w:ascii"), font_name)
        run_properties.rFonts.set(qn("w:hAnsi"), font_name)
        run_properties.rFonts.set(qn("w:eastAsia"), font_name)
        run_properties.rFonts.set(qn("w:cs"), font_name)

    def _get_heading_level(self, section: dict[str, Any]) -> int:
        level = section.get("heading_level", 1)
        return level if isinstance(level, int) and level in (1, 2, 3) else 1

    def _heading_style_name(self, level: int) -> str:
        return {
            1: DOCX_HEADING_1_STYLE_NAME,
            2: DOCX_HEADING_2_STYLE_NAME,
            3: DOCX_HEADING_3_STYLE_NAME,
        }[level]

    def _build_output_path(self, filename: str | None) -> Path:
        """Generate a sanitized, timestamped DOCX filename."""

        if filename is None:
            timestamp = datetime.now().strftime(DOCX_TIMESTAMP_FORMAT)
            filename = f"{GENERATED_DOCUMENT_FILE_PREFIX}_{timestamp}"

        if not isinstance(filename, str) or not filename.strip():
            raise ValueError("DOCX filename must be a non-empty string.")

        filename = self._sanitize_filename(filename.strip())

        if not filename.lower().endswith(DOCX_FILE_EXTENSION):
            filename += DOCX_FILE_EXTENSION
        output_path = self.output_dir / filename

        # Prevent accidental overwrite when a custom name already exists.
        if output_path.exists():
            stem = output_path.stem
            timestamp = datetime.now().strftime(DOCX_TIMESTAMP_FORMAT)
            output_path = self.output_dir / f"{stem}_{timestamp}{DOCX_FILE_EXTENSION}"
        return output_path

    def _sanitize_filename(self, filename: str) -> str:
        filename = re.sub(r'[<>:"/\\|?*]', "_", filename)
        filename = filename.rstrip(". ")

        if not filename:
            raise ValueError("DOCX filename contains no usable characters.")
        return filename

    def _validate_notes(self, notes: dict[str, Any]) -> dict[str, Any]:
        if not isinstance(notes, dict):
            raise TypeError("Structured notes must be a dictionary.")

        title = notes.get("title")
        sections = notes.get("sections")

        if not isinstance(title, str) or not title.strip():
            raise ValueError("Structured notes must contain a non-empty title.")

        if not isinstance(sections, list) or not sections:
            raise ValueError("Structured notes must contain at least one section.")

        for index, section in enumerate(sections, start=1):
            if not isinstance(section, dict):
                raise ValueError(f"Section {index} must be a dictionary.")

            if not isinstance(section.get("heading"), str):
                raise ValueError(f"Section {index} must contain a heading string.")

            for field in (
                "paragraphs",
                "bullets",
                "definitions",
                "examples",
            ):
                if not isinstance(section.get(field, []), list):
                    raise ValueError(f"Section {index} {field} must be a list.")

        if not isinstance(notes.get("summary", []), list):
            raise ValueError("Structured-note summary must be a list.")
        return notes

    def _validate_profile(self, profile: dict[str, Any]) -> dict[str, Any]:
        if not isinstance(profile, dict):
            raise TypeError("Resolved style profile must be a dictionary.")

        missing = [
            section
            for section in self.REQUIRED_PROFILE_SECTIONS
            if not isinstance(profile.get(section), dict)
        ]

        if missing:
            raise ValueError(
                "Resolved style profile is missing these sections: "
                + ", ".join(missing)
            )
        return profile

    def _parse_alignment(self, alignment: Any) -> WD_ALIGN_PARAGRAPH:
        values = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        if not isinstance(alignment, str):
            return WD_ALIGN_PARAGRAPH.LEFT
        return values.get(alignment.strip().lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def _parse_color(self, color: Any) -> RGBColor:
        if not isinstance(color, str):
            color = "#000000"

        color = color.strip().lstrip("#")

        if not re.fullmatch(r"[0-9A-Fa-f]{6}", color):
            color = "000000"
        return RGBColor.from_string(color.upper())

    def _safe_font(self, value: Any) -> str:
        if not isinstance(value, str) or not value.strip():
            return DOCX_FALLBACK_FONT
        return value.strip()

    def _font_size(self, value: Any, fallback: float) -> float:
        return self._safe_number(
            value,
            fallback,
            DOCX_MIN_FONT_SIZE,
            DOCX_MAX_FONT_SIZE,
        )

    def _safe_number(
        self,
        value: Any,
        fallback: float,
        minimum: float,
        maximum: float,
    ) -> float:
        try:
            number = float(value)
        except (TypeError, ValueError):
            return float(fallback)
        return number if minimum <= number <= maximum else float(fallback)

    def _safe_bool(self, value: Any, fallback: bool) -> bool:
        return value if isinstance(value, bool) else fallback

    def _get_spacing(
        self,
        style_data: dict[str, Any],
        primary_key: str,
        fallback_key: str,
    ) -> float:
        value = style_data.get(
            primary_key,
            style_data.get(fallback_key, 0.0),
        )
        return self._safe_number(value, 0.0, 0.0, 72.0)